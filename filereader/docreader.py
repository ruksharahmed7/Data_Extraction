# -*- coding: utf-8 -*-
from docx import Document

import xml.etree.ElementTree as ET
import zipfile as zf
import dataExtraction.fontconvert.fontconverter as fontconverter

import dataExtraction.rulesfile.rules as rules
import traceback



nsprefixes = {
    'mo': 'http://schemas.microsoft.com/office/mac/office/2008/main',
    'o':  'urn:schemas-microsoft-com:office:office',
    've': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    # Text Content
    'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w10': 'urn:schemas-microsoft-com:office:word',
    'wne': 'http://schemas.microsoft.com/office/word/2006/wordml',
    # Drawing
    'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
    'm':   'http://schemas.openxmlformats.org/officeDocument/2006/math',
    'mv':  'urn:schemas-microsoft-com:mac:vml',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'v':   'urn:schemas-microsoft-com:vml',
    'wp':  ('http://schemas.openxmlformats.org/drawingml/2006/wordprocessing'
            'Drawing'),
    # Properties (core and extended)
    'cp':  ('http://schemas.openxmlformats.org/package/2006/metadata/core-pr'
            'operties'),
    'dc':  'http://purl.org/dc/elements/1.1/',
    'ep':  ('http://schemas.openxmlformats.org/officeDocument/2006/extended-'
            'properties'),
    'xsi': 'http://www.w3.org/2001/XMLSchema-instance',
    # Content Types
    'ct':  'http://schemas.openxmlformats.org/package/2006/content-types',
    # Package Relationships
    'r':  ('http://schemas.openxmlformats.org/officeDocument/2006/relationsh'
           'ips'),
    'pr':  'http://schemas.openxmlformats.org/package/2006/relationships',
    # Dublin Core document properties
    'dcmitype': 'http://purl.org/dc/dcmitype/',
    'dcterms':  'http://purl.org/dc/terms/'}

from bijoy2unicode import converter
test = converter.Unicode()
import  re

check_ascii_re=re.compile(r"^[a-zA-Z\d]+$")

def doc_reader_tree_formate(filename):
    document = Document(filename)
    style = document.styles['Normal']
    font = style.font
    font.name = 'SutonnyMJ'
    z = zf.ZipFile(filename)
    f = z.open("word/document.xml")  # a file-like objectclass
    tree = ET.parse(f)  # an ElementTree instance
    paratextlist=[]
    paralist=[]
    for element in tree.iter():
        #print(element.text)
        if element.tag == '{' + nsprefixes['w'] + '}p':
            paralist.append(element)
        if element.tag=="//w:p/w:r/w:br[@w:type='page']" + nsprefixes['w']:
            print('here find:'+element.text)
        # Since a single sentence might be spread over multiple text elements,
        # iterate through each paragraph, appending all text (t) children to that
        # paragraphs text.
    tree_node=0
    para_node=0
    raw_data={}
    converted_data={}
    for para in paralist:
        paratext = ''
        # Loop through each paragraph
        for element in para.iter():
            # Find t (text) elements
            if element.tag == '{' + nsprefixes['w'] + '}t':
                if element.text:
                    paratext = paratext + element.text
                        #toUnicode = test.convertBijoyToUnicode(str(element.text))
                        #paratext = paratext + toUnicode
                    #else:
                     #   toUnicode = test.convertBijoyToUnicode(element.text)
                      #  paratext = paratext + toUnicode
                    #print(paratext.run.font)
            elif element.tag == '{' + nsprefixes['w'] + '}tab':
                paratext = paratext + '\t'
        # Add our completed paragraph text to the list of paragraph text

        if (not len(paratext.strip()) <= 1 and paratext.strip()!=':-' and  paratext.strip()!=':' and  paratext.strip()!='t'):
            data=paratext.strip()
            raw_data[para_node] = data
            #print(data)
            toUnicode=fontconverter.bijoy2uni(data)
            #print(toUnicode)
            paratextlist.append(toUnicode +'^'+str(para_node))
            converted_data[para_node]=toUnicode
        para_node+=1
    #print("Here find all:")
    return paratextlist,raw_data,converted_data

def getParaText(filename):
    doc = Document(filename)

    data = []
    i=0
    idx=0
    for para in doc.paragraphs:
        if(not (rules.clear_re.search(para.text.strip())== None)):
            data.insert(idx,para.text+'{'+str(i)+'}')
            #print(data[idx])
            idx+=1
        i+=1
    return data


def getTableData(filename):
    document = Document(filename)
    # table = document.tables[0]

    data = []

    table_no = 1
    keys = None
    for table in document.tables:
        row_no = 1
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            # print(tuple(text))
            row_data = tuple(text)
            row_data = row_data + (str(table_no), str(row_no))
            data.append(row_data)
            row_no += 1
        table_no += 1
    return data


try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile

WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'

def get_docx_text(path):
            """
            Take the path of a docx file as argument, return the text in unicode.
            """
            document = zipfile.ZipFile(path)
            contentToRead = ["header2.xml", "document.xml", "footer2.xml"]
            paragraphs = []

            for xmlfile in contentToRead:
                xml_content = document.read('word/{}'.format(xmlfile))
                tree = XML(xml_content)
                for paragraph in tree.getiterator(PARA):
                    texts = [node.text
                             for node in paragraph.getiterator(TEXT)
                             if node.text]
                    if texts:
                        textData = ''.join(texts)
                        if xmlfile == "footer2.xml":
                            extractedTxt = "Footer : " + textData
                        elif xmlfile == "header2.xml":
                            extractedTxt = "Header : " + textData
                        else:
                            extractedTxt = textData

                        paragraphs.append(extractedTxt)
            document.close()
            return '\n\n'.join(paragraphs)

def print_doc(filename):
    doc = Document(filename)

    # print the list of paragraphs in the document
    print('List of paragraph objects:->>>')
    print(doc.paragraphs)

    # print the list of the runs
    # in a specified paragraph
    print('\nList of runs objects in 1st paragraph:->>>')
    print(doc.paragraphs[0].runs)

    # print the text in a paragraph
    print('\nText in the 1st paragraph:->>>')
    print(doc.paragraphs[0].text)

    # for printing the complete document
    print('\nThe whole content of the document:->>>\n')
    for i in range(len(doc.paragraphs)):
        #print_doc(doc.paragraphs[i].text)
        for ran in doc.paragraphs[i].runs:
            print(ran.text)
        #print(doc.paragraphs[i].runs)